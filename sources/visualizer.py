import glob
import json
import os
import pickle

# from sources import reader as grdr

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum import style
import docx.styles as styles


# from docx.shared import RGBColor
# from docx.shared import Pt

class Visualizer:
    """
    Reads in the log files and extracts the necessary information needed
    to a docx (pdf later) file

    Arguments:
        - path_logs: the path to the logs
        - mturk_session: the mturk dev session

    Returns:
        - A word file which shows the necessary information from the logs
    """

    def __init__(self, path_logs, mturk_session):
        self.path_results = os.path.join(path_logs + "/output/", mturk_session)
        self.path_logs = os.path.join(path_logs, mturk_session)

    def run(self, global_rdr):
        # This if condition is for using the imported reader class
        # Differentiates between when called from here or from the run.py file (change style later)
        if not global_rdr:
            logs_reader = rdr.ReadLogs(self.path_logs)
        else:
            logs_reader = global_rdr.ReadLogs(self.path_logs)
            # logs_reader = grdr.ReadLogs(self.path_logs, self.mturk_session)

        # This returns a list of the speakers userid in each meetup room and the corresponding list of json logtexts
        list_users, list_logtext = logs_reader.run()
        for _users, single_logtext in zip(self.list_users, self.list_logtext):
            self._write_to_docx(_users, single_logtext)
        print("DONE")

    def _write_to_docx(self, _users, single_logtext, includeProcessed=False, single_sample=None):

        """
        This function creates a document (word file) for each logfile and writes the
        necessary information into it (pdf later)
        """
        ##########################################################
        #        Create a dict for structuring the inputs        #
        ##########################################################
        dict_inputs = {'table_traits': ['Traits / Attitudes', 1, 1, 'Persona'],
                       'table_facts': ['Facts / Knowledge', 1, 2, 'Entity', 'Content'],
                       'table_questions': ['Questions', 1, 1, ''],
                       'table_answers': ['Answers', 1, 2, 'Entity', 'Answer']}

        story = None
        for item in single_logtext:
            if item['user']['name'] == "Moderator" and item['type'] == "story_type":
                story = item['story_type']

            if item['type'] == "table_entities":
                table_entity = item['table']

                # Create document for saving the conversation and other factors
        # Create a new document for each meetup log
        self.document = Document()
        if story is None:
            self.document.add_heading("Story type not specified", 0)
        else:
            self.document.add_heading(story, 0)

        curr_userid = 0
        dialogue_list = []
        dialogue_tok_list = []

        for item in single_logtext:
            ######################################################################################
            #                       Input necessary info using dict_inputs                       #
            ######################################################################################
            # Write the facts and traits for both users
            for _id in _users:
                if _id == _users[0]:
                    speaker = 'First speaker'
                else:
                    speaker = 'Second speaker'

                for input, params in dict_inputs.items():
                    if item['type'] == input and item['user']['id'] == _id:
                        if input == 'table_traits':  # Only add this heading at the first table of each speaker
                            p = self.document.add_heading(speaker, level=1)
                            p.add_run().bold = True
                            # p.add_run('bold').bold = True
                            # p.add_run('italic.').italic = True

                        p = self.document.add_heading(params[0], level=2)
                        p.add_run().italic = True

                        if len(params) > 4:  # e.g for case: table_facts and table_answers
                            table = self.document.add_table(rows=params[1], cols=params[2])
                            # table.style = 'Light Grid Accent 1'
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = params[3]
                            hdr_cells[1].text = params[4]

                            # Input facts correspondingly (key = topic, value = fact)
                            # Note: zip only works if the two iterations are of same length
                            for _entity, _content in zip(item['table'][0], item['table'][1]):
                                row_cells = table.add_row().cells
                                row_cells[0].text = _entity
                                row_cells[1].text = _content

                        else:  # e.g for case: table_attributes and table_questions
                            for _content in item['table']:
                                if input == 'table_traits':
                                    self.document.add_paragraph(_content, style='List')
                                else:
                                    self.document.add_paragraph(_content, style='List Bullet')

            ######################################################
            #        Rearrange the dialogue speaker turns        #
            ######################################################
            # Rearrange the dialogue orders so that its just one line per speaker turn
            if item['type'] == 'text' and item['user']['id'] in _users:
                if item['user']['id'] == curr_userid:
                    # Update the dialogue_list if the same speaker has used two turns
                    msg = tmp_msg + " [EOU] " + item['msg']
                    dialogue_list[len(dialogue_list) - 1] = msg
                    tmp_msg = msg
                else:
                    dialogue_list.append(item['msg'])
                    # tmp_msg is used for when a speaker uses two turns (to put each of them together)
                    tmp_msg = item['msg']
                curr_userid = item['user']['id']

            if item['type'] == "join" and item['user']['name'] == "Moderator":
                uid = item['timestamp-iso'] + "-" + item['room']['name']
                # use for naming the file
                # this replace is performed because files cant be saved with character ":"
                fname = uid.replace(":", "-")
                fname = fname + ".docx"
        # --------------------------------    END OF FOR LOOP    --------------------------------#

        #####################################################
        #            Input entities provided                #
        #####################################################
        e = self.document.add_heading('Entities Provided To Users', level=1)
        e.add_run().bold = True

        for ent in table_entity:
            self.document.add_paragraph(ent, style='List Bullet')

        ################################################
        #            Input the dialogue                #
        ################################################
        p = self.document.add_heading('Dialogue', level=1)
        p.add_run().bold = True

        for y in range(0, len(dialogue_list)):
            if y % 2 == 0:  # Even for First speaker
                self.document.add_paragraph(dialogue_list[y])
            else:  # Odd numbers for Second speaker
                paragraph = self.document.add_paragraph(dialogue_list[y])
                # Make this paragraph to align right
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if includeProcessed:
            self._include_processed_dialogue(single_sample)
        ######################
        #        Save        #
        ######################
        self._saver(self.document, fname)

        print("\nVisualization result has been saved into {}".format(self.path_results))
        return True

    def _include_processed_dialogue(self, single_sample):
        """
        Create a document (word file) for the processed logfile

        single_sample_output.update({
            'uid': uid,
            'story': story,
            'entity_provided': table_entity,
            'knowledge': knowledge_dict,
            'attitudes': attitudes_dict,
            'questions': question_dict,
            'answers': answer_dict,
            'named_entities': 'empty',
            'dialogue_orig': self.dialogue,
            'dialogue_tokenized': dialogue_tokenized,
            'dialogue_processesd': 'empty',
            'dialogue_named_entity': 'empty'
            })
        """
        #########################################
        #        Add the other dialogues        #
        #########################################
        dict_inputs = {'dialogue_original': 'Original Dialogue',
                       'dialogue_processesd': 'Processesd Dialogue',
                       'dialogue_named_entity': 'Named Entity Dialogue'}
        ###################################################################
        #            Input the selected processed dialogue                #
        ###################################################################
        for entity in single_sample:
            if entity in dict_inputs:
                p = self.document.add_heading(dict_inputs[entity], level=1)
                p.add_run().bold = True

                if single_sample[entity] == 'empty':
                    self.document.add_paragraph("{} not yet added".format(entity))
                else:
                    dialogue_list = single_sample[entity]
                    for y in range(0, len(dialogue_list)):
                        if y % 2 == 0:  # Even for First speaker
                            self.document.add_paragraph(dialogue_list[y])
                        else:  # Odd numbers for Second speaker
                            paragraph = self.document.add_paragraph(dialogue_list[y])
                            # Make this paragraph to align right
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _saver(self, _document, _fname):
        if not os.path.exists(self.path_results):
            os.makedirs(self.path_results)

        doc_path = os.path.join(self.path_results, _fname)
        _document.save(doc_path)


if __name__ == '__main__':
    import reader as rdr

    # mturk_session = "moviedatagen_real_test_2/"
    mturk_session = "data_collection_part_1/"

    processed_log_path = ""
    vis = Visualizer(path_logs="../logfiles/", mturk_session=mturk_session)
    vis.run("")
